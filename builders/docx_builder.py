from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import json
from pathlib import Path
from docx2pdf import convert
from schema.schema import (
    Resume,
    Header,
    Education,
    Experience,
    Project,
    TechnicalSkills
)

class _Docx_Builder():
    def __init__(self):
        pass

    def add_run(self, paragraph, text, bold=False, italic=False, size=9.5, underline=False):
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline
        run.font.size = Pt(size)
        return run

    def add_hyperlink(self, paragraph: Paragraph, text: str, url: str):
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run_elem = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        style = OxmlElement("w:rStyle")
        style.set(qn("w:val"), "Hyperlink")
        rpr.append(style)
        run_elem.append(rpr)
        t = OxmlElement("w:t")
        t.text = text
        run_elem.append(t)
        hyperlink.append(run_elem)
        paragraph._p.append(hyperlink)

    def section_heading(self, doc: Document, label):
        p = doc.add_paragraph()
        run = p.add_run(label.upper())
        run.bold = True
        run.font.size = Pt(10)

        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

        p.paragraph_format.space_before = Pt(6)     # slightly more breathing room
        p.paragraph_format.space_after = Pt(2)
        return p

    def sub_heading(self, doc, title, date):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)

        self.add_run(p, title, bold=True, size=9.5)
        self.add_run(p, "\t", size=9.5)
        self.add_run(p, date, italic=True, size=9.5)

        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), "9360")
        tabs.append(tab)
        pPr.append(tabs)
        return p

    def org_line(self, doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        self.add_run(p, text, italic=True, underline=True, size=9)
        return p

    def bullet_point(self, doc, text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)      # tiny gap between bullets — not 0
        p.paragraph_format.left_indent = Inches(0.25)
        self.add_run(p, text, size=9)
        return p

    def skill_line(self, doc, label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        self.add_run(p, label, bold=True, size=9.5)
        self.add_run(p, " " + value, size=9.5)
        return p


class Resume_Builder(_Docx_Builder):
    def __init__(self, json_data: str, ):
        # self.filename = filename
        self.doc = Document()
        self.section = self.doc.sections[0]
        self.section.page_width = Inches(8.5)
        self.section.page_height = Inches(11)
        self.section.top_margin = Inches(0.5)       # back to normal
        self.section.bottom_margin = Inches(0.5)    # back to normal
        self.section.left_margin = Inches(0.55)
        self.section.right_margin = Inches(0.55)

        self.doc.styles["Normal"].font.name = "Calibri"
        self.doc.styles["Normal"].font.size = Pt(9.5)

        # comfortable line spacing — not too tight
        normal_style = self.doc.styles["Normal"]
        pPr = normal_style.element.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), "236")        # between tight(220) and single(240)
        spacing.set(qn("w:lineRule"), "auto")
        pPr.append(spacing)

        self.resume_data = json_data
        self.create_doc(data=self.resume_data )

    def _header(self, header: Header):
        name_p = self.doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_run(name_p, text=header["full_name"], bold=True, size=15)
        name_p.paragraph_format.space_after = Pt(2)

        contact_p = self.doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_run(contact_p, f"{header['email']} | {header['number']} | {header['location']} | ", size=9)
        self.add_hyperlink(contact_p, f"{header['linkdin']}", f"https://{header['linkdin']}")
        self.add_run(contact_p, " | ", size=9)
        self.add_hyperlink(contact_p, f"{header['github']}", f"https://{header['github']}")
        contact_p.paragraph_format.space_after = Inches(0.25)

    def _summary(self, summary: str):
        self.section_heading(self.doc, "Summary")
        p = self.doc.add_paragraph()
        self.add_run(p, text=summary, size=9)
        p.paragraph_format.space_after = Pt(2)

    def _education(self, education: Education):
        self.section_heading(self.doc, label="Education")
        self.sub_heading(self.doc, f"{education['school']} - {education['study']}", f"Expected: {education['expected']}")
        self.bullet_point(self.doc, text=f"GPA: {education['GPA']} | {education['award']}")
        # limit courses to avoid wrapping to 3 lines
        courses = education['Relevant Courses'][:6]
        self.bullet_point(self.doc, f"Relevant coursework: {', '.join(courses)}")

    def _experience(self, experience: Experience):
        self.section_heading(self.doc, "Work Experience")
        for job in experience:
            # skip jobs with no title or empty proposed entries
            if not job.get("timeline"):
                continue
            self.sub_heading(self.doc, job["job_title"], job["timeline"])
            if job["organization"]:
                self.org_line(self.doc, f"{job['organization']} | {job['location']}")
            for task in job["responsiblities"]:
                self.bullet_point(self.doc, task)

    def _projects(self, projects: Project):
        self.section_heading(self.doc, "Personal Projects")
        for project in projects:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            self.add_run(p, f"{project['name']} — {project['description']} (", bold=True, size=9)
            self.add_hyperlink(p, "github", f"https://{project['link']}")
            self.add_run(p, ")", bold=True, size=9)
            for point in project["bullet_points"]:
                self.bullet_point(self.doc, point)

    def _technical_skills(self, technical_skills: TechnicalSkills):
        self.section_heading(self.doc, "Technical Skills")
        for key, val in technical_skills.items():
            self.skill_line(self.doc, label=f"{key}:", value=f"{', '.join(val)}")

    def _export_resume(self, filename: str = "resume", output_dir: str = "resume_ark/", keep_docx: bool = False):
        docx_path = Path(output_dir) / f"{filename}.docx"
        pdf_path = Path(output_dir) / f"{filename}.pdf"
        self.doc.save(docx_path)
        convert(docx_path, pdf_path)
        if not keep_docx:
            docx_path.unlink()
        print(f"Saved: {pdf_path}")

    def create_doc(self, data: Resume):
        self._header(data["header"])
        self._summary(data["summary"])
        self._education(data["education"])
        self._experience(data["experience"])
        self._projects(data["projects"])
        self._technical_skills(data["technical_skills"])
        self._export_resume()
        

if __name__ == "__main__":
  
    json_data_path = "./schema/resume_schema.json"
    data = {}
    if Path(json_data_path).exists():
        with open(json_data_path, "r") as file:
            data = json.load(file)
        Resume_Builder(data)
            
    else:
        print("couldn't find json file")
    


    
    